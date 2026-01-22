import os
import io
import json
import zipfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import List, Optional

import httpx
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.requests import Request
from dotenv import load_dotenv
from pydantic import BaseModel

load_dotenv()

app = FastAPI(title="Экспертиза локальных актов РГГУ")

BASE_DIR = Path(__file__).parent
UPLOADS_DIR = BASE_DIR / "uploads"
EXPORTS_DIR = BASE_DIR / "exports"
TEMPLATES_DIR = BASE_DIR / "templates"

UPLOADS_DIR.mkdir(exist_ok=True)
EXPORTS_DIR.mkdir(exist_ok=True)

templates = Jinja2Templates(directory=str(TEMPLATES_DIR))

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

# Хранилище результатов в памяти
results_store: dict = {}


# ==================== НОРМАТИВНЫЕ АКТЫ ====================

UKAZ_809_TEXT = """
УКАЗ ПРЕЗИДЕНТА РОССИЙСКОЙ ФЕДЕРАЦИИ ОТ 09.11.2022 № 809
"Об утверждении Основ государственной политики по сохранению и укреплению
традиционных российских духовно-нравственных ценностей"

КЛЮЧЕВЫЕ ПОЛОЖЕНИЯ:

1. Традиционные ценности - это нравственные ориентиры, формирующие мировоззрение
граждан России, передаваемые от поколения к поколению, лежащие в основе
общероссийской гражданской идентичности и единого культурного пространства страны.

2. К традиционным ценностям относятся:
   - жизнь, достоинство, права и свободы человека
   - патриотизм, гражданственность, служение Отечеству
   - высокие нравственные идеалы
   - крепкая семья, созидательный труд
   - приоритет духовного над материальным
   - гуманизм, милосердие, справедливость
   - коллективизм, взаимопомощь и взаимоуважение
   - историческая память и преемственность поколений
   - единство народов России

3. Угрозы традиционным ценностям:
   - деятельность экстремистских и террористических организаций
   - действия США и других недружественных государств
   - деструктивная идеология
   - деятельность транснациональных корпораций и иностранных НКО

4. Цели государственной политики:
   - сохранение и укрепление традиционных ценностей
   - передача их от поколения к поколению
   - противодействие распространению деструктивной идеологии
   - формирование на международной арене образа России как хранителя ценностей

5. Задачи в сфере образования и воспитания:
   - воспитание в духе уважения к традиционным ценностям
   - поддержка общественных проектов в сфере патриотического воспитания
   - сохранение исторической памяти, противодействие фальсификации истории
   - поддержка религиозных организаций традиционных конфессий
   - защита института семьи, материнства, отцовства и детства
"""

RASPORYAZHENIE_1734_TEXT = """
РАСПОРЯЖЕНИЕ ПРАВИТЕЛЬСТВА РОССИЙСКОЙ ФЕДЕРАЦИИ ОТ 01.07.2024 № 1734-р
"О Плане мероприятий по реализации в 2024-2026 г.г. Основ государственной политики
по сохранению и укреплению традиционных российских духовно-нравственных ценностей"

КЛЮЧЕВЫЕ МЕРОПРИЯТИЯ ПЛАНА:

1. Нормативно-правовое обеспечение:
   - Подготовка предложений по совершенствованию законодательства
   - Мониторинг реализации государственной политики в сфере ценностей

2. Образование и воспитание:
   - Разработка методических рекомендаций для образовательных организаций
   - Включение вопросов традиционных ценностей в образовательные программы
   - Проведение всероссийских мероприятий патриотической направленности
   - Поддержка добровольческих движений

3. Культура и искусство:
   - Поддержка проектов, направленных на сохранение традиционных ценностей
   - Создание контента, продвигающего традиционные ценности
   - Противодействие деструктивному контенту

4. Информационная политика:
   - Формирование позитивного информационного пространства
   - Защита детей от деструктивной информации
   - Развитие медиаграмотности

5. Работа с молодёжью:
   - Поддержка молодёжных организаций патриотической направленности
   - Развитие наставничества
   - Вовлечение молодёжи в общественно полезную деятельность

6. Международное сотрудничество:
   - Продвижение традиционных ценностей на международных площадках
   - Поддержка соотечественников за рубежом

7. Научное обеспечение:
   - Проведение исследований в сфере традиционных ценностей
   - Мониторинг угроз традиционным ценностям
"""


# ==================== ПРОМПТ ДЛЯ АНАЛИЗА ====================

ANALYSIS_PROMPT = """Ты — эксперт по правовому анализу локальных нормативных актов образовательных организаций.

Твоя задача: проанализировать предоставленный документ (локальный нормативный акт РГГУ) на предмет соответствия двум федеральным документам:

1. Указ Президента РФ от 09.11.2022 № 809 "Об утверждении Основ государственной политики по сохранению и укреплению традиционных российских духовно-нравственных ценностей"

2. Распоряжение Правительства РФ от 01.07.2024 № 1734-р "О Плане мероприятий по реализации в 2024-2026 г.г. Основ государственной политики по сохранению и укреплению традиционных российских духовно-нравственных ценностей"

=== СОДЕРЖАНИЕ УКАЗА № 809 ===
{ukaz_text}

=== СОДЕРЖАНИЕ РАСПОРЯЖЕНИЯ № 1734-р ===
{rasporyazhenie_text}

=== АНАЛИЗИРУЕМЫЙ ДОКУМЕНТ ===
Название: {doc_name}

Текст документа:
{doc_text}

=== ИНСТРУКЦИИ ===

Проведи детальный анализ и верни результат СТРОГО в следующем JSON формате:

{{
    "document_name": "название анализируемого документа",
    "summary": "краткое описание содержания документа (2-3 предложения)",

    "ukaz_809": {{
        "matches": [
            {{
                "doc_reference": "ссылка на пункт/раздел анализируемого документа",
                "ukaz_reference": "ссылка на положение Указа № 809",
                "description": "описание соответствия"
            }}
        ],
        "contradictions": [
            {{
                "doc_reference": "ссылка на пункт/раздел документа или 'отсутствует'",
                "ukaz_reference": "ссылка на положение Указа № 809",
                "description": "описание расхождения или отсутствующего элемента",
                "recommendation": "рекомендация по устранению"
            }}
        ]
    }},

    "rasporyazhenie_1734": {{
        "matches": [
            {{
                "doc_reference": "ссылка на пункт/раздел анализируемого документа",
                "rasp_reference": "ссылка на пункт Плана мероприятий",
                "description": "описание соответствия"
            }}
        ],
        "contradictions": [
            {{
                "doc_reference": "ссылка на пункт/раздел документа или 'отсутствует'",
                "rasp_reference": "ссылка на пункт Плана мероприятий",
                "description": "описание расхождения или отсутствующего элемента",
                "recommendation": "рекомендация по устранению"
            }}
        ]
    }},

    "conclusion": {{
        "status": "соответствует" | "частично соответствует" | "требует доработки",
        "summary": "итоговое заключение (2-3 предложения)",
        "priority_recommendations": ["список приоритетных рекомендаций по доработке"]
    }}
}}

ВАЖНО:
- Если соответствий или противоречий нет, оставь пустой массив []
- Будь конкретен в ссылках на пункты документов
- Учитывай специфику образовательной организации
- Отвечай ТОЛЬКО валидным JSON без дополнительного текста
"""


# ==================== ФУНКЦИИ ПАРСИНГА ====================

def extract_text_from_pdf(file_path: Path) -> str:
    """Извлечь текст из PDF файла."""
    try:
        doc = fitz.open(str(file_path))
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Ошибка чтения PDF: {str(e)}")


def extract_text_from_docx(file_path: Path) -> str:
    """Извлечь текст из DOCX файла."""
    try:
        if not file_path.exists():
            raise Exception(f"Файл не найден: {file_path}")
        print(f"DEBUG: Reading DOCX from {file_path}")
        doc = Document(str(file_path))
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Ошибка чтения DOCX: {str(e)}")


def extract_text_from_doc(file_path: Path) -> str:
    """Извлечь текст из старого DOC файла (Word 97-2003) через textutil на macOS."""
    import subprocess
    import tempfile

    try:
        # Проверяем существование файла
        if not file_path.exists():
            raise Exception(f"Файл не найден: {file_path}")

        # Создаём временный файл для вывода
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
            tmp_path = tmp.name

        # Конвертируем через textutil (macOS)
        result = subprocess.run(
            ['textutil', '-convert', 'txt', '-output', tmp_path, str(file_path)],
            capture_output=True,
            text=True
        )

        if result.returncode != 0:
            raise Exception(f"textutil error: {result.stderr}")

        # Читаем результат
        text = Path(tmp_path).read_text(encoding="utf-8")

        # Удаляем временный файл
        Path(tmp_path).unlink(missing_ok=True)

        return text.strip()
    except FileNotFoundError:
        raise HTTPException(status_code=400, detail="Формат .doc не поддерживается на этой системе (требуется macOS)")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Ошибка чтения DOC: {str(e)}")


def extract_text(file_path: Path) -> str:
    """Извлечь текст из файла в зависимости от формата."""
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        return extract_text_from_docx(file_path)
    elif suffix == ".doc":
        return extract_text_from_doc(file_path)
    elif suffix == ".txt":
        return file_path.read_text(encoding="utf-8")
    else:
        raise HTTPException(status_code=400, detail=f"Неподдерживаемый формат: {suffix}")


# ==================== LLM ИНТЕГРАЦИЯ ====================

async def analyze_document(doc_name: str, doc_text: str) -> dict:
    """Отправить документ на анализ через OpenRouter API."""

    prompt = ANALYSIS_PROMPT.format(
        ukaz_text=UKAZ_809_TEXT,
        rasporyazhenie_text=RASPORYAZHENIE_1734_TEXT,
        doc_name=doc_name,
        doc_text=doc_text[:50000]  # Ограничение по длине
    )

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:8000",
        "X-Title": "RGGU Expertise Service"
    }

    payload = {
        "model": "anthropic/claude-3.5-sonnet",
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 4096,
        "temperature": 0.1
    }

    async with httpx.AsyncClient(timeout=120.0) as client:
        response = await client.post(OPENROUTER_URL, json=payload, headers=headers)

        if response.status_code != 200:
            raise HTTPException(
                status_code=500,
                detail=f"Ошибка API: {response.status_code} - {response.text}"
            )

        result = response.json()
        content = result["choices"][0]["message"]["content"]

        # Парсим JSON из ответа
        try:
            # Убираем возможные markdown блоки
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0]
            elif "```" in content:
                content = content.split("```")[1].split("```")[0]

            return json.loads(content.strip())
        except json.JSONDecodeError as e:
            # Если не удалось распарсить, возвращаем сырой текст
            return {
                "document_name": doc_name,
                "raw_response": content,
                "parse_error": str(e)
            }


# ==================== ГЕНЕРАЦИЯ DOCX ====================

def create_expertise_docx(analysis: dict) -> io.BytesIO:
    """Создать DOCX файл с экспертным заключением."""

    doc = Document()

    # Заголовок
    title = doc.add_heading("ЭКСПЕРТНОЕ ЗАКЛЮЧЕНИЕ", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Информация о документе
    doc.add_paragraph()
    doc.add_paragraph(f"Наименование документа: {analysis.get('document_name', 'Не указано')}")
    doc.add_paragraph(f"Дата экспертизы: {datetime.now().strftime('%d.%m.%Y')}")

    if analysis.get('summary'):
        doc.add_paragraph()
        doc.add_paragraph(f"Краткое содержание: {analysis['summary']}")

    # Анализ по Указу № 809
    doc.add_heading("АНАЛИЗ НА СООТВЕТСТВИЕ УКАЗУ ПРЕЗИДЕНТА РФ ОТ 09.11.2022 № 809", level=1)

    ukaz = analysis.get('ukaz_809', {})

    # Соответствия
    doc.add_heading("Соответствует:", level=2)
    matches = ukaz.get('matches', [])
    if matches:
        for m in matches:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{m.get('doc_reference', '')}: ").bold = True
            p.add_run(f"{m.get('description', '')} ")
            p.add_run(f"(соотв. {m.get('ukaz_reference', '')})").italic = True
    else:
        doc.add_paragraph("Явных соответствий не выявлено.", style='List Bullet')

    # Расхождения
    doc.add_heading("Расходится:", level=2)
    contradictions = ukaz.get('contradictions', [])
    if contradictions:
        for c in contradictions:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{c.get('doc_reference', '')}: ").bold = True
            p.add_run(f"{c.get('description', '')} ")
            p.add_run(f"(см. {c.get('ukaz_reference', '')})").italic = True
            if c.get('recommendation'):
                rec = doc.add_paragraph(style='List Bullet 2')
                rec.add_run("Рекомендация: ").bold = True
                rec.add_run(c['recommendation'])
    else:
        doc.add_paragraph("Расхождений не выявлено.", style='List Bullet')

    # Анализ по Распоряжению № 1734-р
    doc.add_heading("АНАЛИЗ НА СООТВЕТСТВИЕ РАСПОРЯЖЕНИЮ ПРАВИТЕЛЬСТВА РФ ОТ 01.07.2024 № 1734-р", level=1)

    rasp = analysis.get('rasporyazhenie_1734', {})

    # Соответствия
    doc.add_heading("Соответствует:", level=2)
    matches = rasp.get('matches', [])
    if matches:
        for m in matches:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{m.get('doc_reference', '')}: ").bold = True
            p.add_run(f"{m.get('description', '')} ")
            p.add_run(f"(соотв. {m.get('rasp_reference', '')})").italic = True
    else:
        doc.add_paragraph("Явных соответствий не выявлено.", style='List Bullet')

    # Расхождения
    doc.add_heading("Расходится:", level=2)
    contradictions = rasp.get('contradictions', [])
    if contradictions:
        for c in contradictions:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{c.get('doc_reference', '')}: ").bold = True
            p.add_run(f"{c.get('description', '')} ")
            p.add_run(f"(см. {c.get('rasp_reference', '')})").italic = True
            if c.get('recommendation'):
                rec = doc.add_paragraph(style='List Bullet 2')
                rec.add_run("Рекомендация: ").bold = True
                rec.add_run(c['recommendation'])
    else:
        doc.add_paragraph("Расхождений не выявлено.", style='List Bullet')

    # Итоговое заключение
    doc.add_heading("ИТОГОВОЕ ЗАКЛЮЧЕНИЕ", level=1)

    conclusion = analysis.get('conclusion', {})
    status = conclusion.get('status', 'не определён')

    p = doc.add_paragraph()
    p.add_run("Статус: ").bold = True
    p.add_run(status.upper())

    if conclusion.get('summary'):
        doc.add_paragraph(conclusion['summary'])

    if conclusion.get('priority_recommendations'):
        doc.add_paragraph()
        doc.add_paragraph("Приоритетные рекомендации:").bold = True
        for rec in conclusion['priority_recommendations']:
            doc.add_paragraph(rec, style='List Number')

    # Сохраняем в BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==================== API ЭНДПОИНТЫ ====================

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    """Главная страница."""
    return templates.TemplateResponse("index.html", {"request": request})


class AnalysisResult(BaseModel):
    session_id: str
    results: List[dict]


@app.post("/api/analyze")
async def analyze_files(files: List[UploadFile] = File(...)):
    """Загрузить и проанализировать документы."""

    if len(files) > 20:
        raise HTTPException(status_code=400, detail="Максимум 20 файлов за раз")

    session_id = str(uuid.uuid4())
    results = []

    for file in files:
        # Сохраняем файл с безопасным именем (без пробелов и спецсимволов)
        import re
        safe_filename = re.sub(r'[^\w\.\-]', '_', file.filename)
        file_path = UPLOADS_DIR / f"{session_id}_{safe_filename}"
        content = await file.read()
        file_path.write_bytes(content)
        print(f"DEBUG: Saved file to {file_path}, exists: {file_path.exists()}, size: {len(content)}")

        try:
            # Извлекаем текст
            text = extract_text(file_path)

            if not text.strip():
                results.append({
                    "filename": file.filename,
                    "error": "Не удалось извлечь текст из документа"
                })
                continue

            # Анализируем
            analysis = await analyze_document(file.filename, text)
            analysis["filename"] = file.filename
            results.append(analysis)

        except Exception as e:
            results.append({
                "filename": file.filename,
                "error": str(e)
            })
        finally:
            # Удаляем временный файл
            if file_path.exists():
                file_path.unlink()

    # Сохраняем результаты
    results_store[session_id] = results

    return {"session_id": session_id, "results": results}


@app.get("/api/export/{session_id}/{index}")
async def export_single(session_id: str, index: int):
    """Скачать экспертизу по одному документу в DOCX."""

    if session_id not in results_store:
        raise HTTPException(status_code=404, detail="Сессия не найдена")

    results = results_store[session_id]

    if index < 0 or index >= len(results):
        raise HTTPException(status_code=404, detail="Документ не найден")

    analysis = results[index]

    if "error" in analysis:
        raise HTTPException(status_code=400, detail="Для этого документа нет экспертизы")

    from urllib.parse import quote
    docx_buffer = create_expertise_docx(analysis)
    original_name = analysis.get('filename', 'document')
    # Убираем расширение если есть
    base_name = original_name.rsplit('.', 1)[0] if '.' in original_name else original_name
    # ASCII имя - только латиница и цифры
    filename_ascii = f"Expertiza_{index}.docx"
    filename_utf8 = quote(f"Экспертиза_{base_name}.docx")

    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"{filename_ascii}\"; filename*=UTF-8''{filename_utf8}"}
    )


@app.get("/api/export-all/{session_id}")
async def export_all(session_id: str):
    """Скачать все экспертизы в ZIP архиве."""

    if session_id not in results_store:
        raise HTTPException(status_code=404, detail="Сессия не найдена")

    results = results_store[session_id]

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for i, analysis in enumerate(results):
            if "error" not in analysis:
                docx_buffer = create_expertise_docx(analysis)
                filename = f"Экспертиза_{analysis.get('filename', f'документ_{i}')}.docx"
                zf.writestr(filename, docx_buffer.getvalue())

    zip_buffer.seek(0)

    from urllib.parse import quote
    filename = f"Expertiza_RGGU_{session_id[:8]}.zip"
    filename_utf8 = quote(f"Экспертизы_РГГУ_{session_id[:8]}.zip")

    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={filename}; filename*=UTF-8''{filename_utf8}"}
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
