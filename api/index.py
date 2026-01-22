import os
import io
import json
import zipfile
import uuid
import re
from datetime import datetime
from pathlib import Path
from typing import List
from urllib.parse import quote

import httpx
import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

app = FastAPI(title="Экспертиза локальных актов РГГУ")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Vercel использует /tmp для временных файлов
UPLOADS_DIR = Path("/tmp/uploads")
UPLOADS_DIR.mkdir(exist_ok=True)

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

results_store: dict = {}

# ==================== НОРМАТИВНЫЕ АКТЫ ====================

UKAZ_809_TEXT = """
УКАЗ ПРЕЗИДЕНТА РОССИЙСКОЙ ФЕДЕРАЦИИ ОТ 09.11.2022 № 809
"Об утверждении Основ государственной политики по сохранению и укреплению
традиционных российских духовно-нравственных ценностей"

КЛЮЧЕВЫЕ ПОЛОЖЕНИЯ:

1. Традиционные ценности - это нравственные ориентиры, формирующие мировоззрение
граждан России, передаваемые от поколения к поколению, лежащие в основе
общероссийской гражданской идентичности и единого культурного пространства страны.

2. К традиционным ценностям относятся:
   - жизнь, достоинство, права и свободы человека
   - патриотизм, гражданственность, служение Отечеству
   - высокие нравственные идеалы
   - крепкая семья, созидательный труд
   - приоритет духовного над материальным
   - гуманизм, милосердие, справедливость
   - коллективизм, взаимопомощь и взаимоуважение
   - историческая память и преемственность поколений
   - единство народов России

3. Угрозы традиционным ценностям:
   - деятельность экстремистских и террористических организаций
   - действия США и других недружественных государств
   - деструктивная идеология
   - деятельность транснациональных корпораций и иностранных НКО

4. Цели государственной политики:
   - сохранение и укрепление традиционных ценностей
   - передача их от поколения к поколению
   - противодействие распространению деструктивной идеологии
   - формирование на международной арене образа России как хранителя ценностей

5. Задачи в сфере образования и воспитания:
   - воспитание в духе уважения к традиционным ценностям
   - поддержка общественных проектов в сфере патриотического воспитания
   - сохранение исторической памяти, противодействие фальсификации истории
   - поддержка религиозных организаций традиционных конфессий
   - защита института семьи, материнства, отцовства и детства
"""

RASPORYAZHENIE_1734_TEXT = """
РАСПОРЯЖЕНИЕ ПРАВИТЕЛЬСТВА РОССИЙСКОЙ ФЕДЕРАЦИИ ОТ 01.07.2024 № 1734-р
"О Плане мероприятий по реализации в 2024-2026 г.г. Основ государственной политики
по сохранению и укреплению традиционных российских духовно-нравственных ценностей"

КЛЮЧЕВЫЕ МЕРОПРИЯТИЯ ПЛАНА:

1. Нормативно-правовое обеспечение:
   - Подготовка предложений по совершенствованию законодательства
   - Мониторинг реализации государственной политики в сфере ценностей

2. Образование и воспитание:
   - Разработка методических рекомендаций для образовательных организаций
   - Включение вопросов традиционных ценностей в образовательные программы
   - Проведение всероссийских мероприятий патриотической направленности
   - Поддержка добровольческих движений

3. Культура и искусство:
   - Поддержка проектов, направленных на сохранение традиционных ценностей
   - Создание контента, продвигающего традиционные ценности
   - Противодействие деструктивному контенту

4. Информационная политика:
   - Формирование позитивного информационного пространства
   - Защита детей от деструктивной информации
   - Развитие медиаграмотности

5. Работа с молодёжью:
   - Поддержка молодёжных организаций патриотической направленности
   - Развитие наставничества
   - Вовлечение молодёжи в общественно полезную деятельность

6. Международное сотрудничество:
   - Продвижение традиционных ценностей на международных площадках
   - Поддержка соотечественников за рубежом

7. Научное обеспечение:
   - Проведение исследований в сфере традиционных ценностей
   - Мониторинг угроз традиционным ценностям
"""

ANALYSIS_PROMPT = """Ты — эксперт по правовому анализу локальных нормативных актов образовательных организаций.

Твоя задача: проанализировать предоставленный документ (локальный нормативный акт РГГУ) на предмет соответствия двум федеральным документам:

1. Указ Президента РФ от 09.11.2022 № 809 "Об утверждении Основ государственной политики по сохранению и укреплению традиционных российских духовно-нравственных ценностей"

2. Распоряжение Правительства РФ от 01.07.2024 № 1734-р "О Плане мероприятий по реализации в 2024-2026 г.г. Основ государственной политики по сохранению и укреплению традиционных российских духовно-нравственных ценностей"

=== СОДЕРЖАНИЕ УКАЗА № 809 ===
{ukaz_text}

=== СОДЕРЖАНИЕ РАСПОРЯЖЕНИЯ № 1734-р ===
{rasporyazhenie_text}

=== АНАЛИЗИРУЕМЫЙ ДОКУМЕНТ ===
Название: {doc_name}

Текст документа:
{doc_text}

=== ИНСТРУКЦИИ ===

Проведи детальный анализ и верни результат СТРОГО в следующем JSON формате:

{{
    "document_name": "название анализируемого документа",
    "summary": "краткое описание содержания документа (2-3 предложения)",

    "ukaz_809": {{
        "matches": [
            {{
                "doc_reference": "ссылка на пункт/раздел анализируемого документа",
                "ukaz_reference": "ссылка на положение Указа № 809",
                "description": "описание соответствия"
            }}
        ],
        "contradictions": [
            {{
                "doc_reference": "ссылка на пункт/раздел документа или 'отсутствует'",
                "ukaz_reference": "ссылка на положение Указа № 809",
                "description": "описание расхождения или отсутствующего элемента",
                "recommendation": "рекомендация по устранению"
            }}
        ]
    }},

    "rasporyazhenie_1734": {{
        "matches": [
            {{
                "doc_reference": "ссылка на пункт/раздел анализируемого документа",
                "rasp_reference": "ссылка на пункт Плана мероприятий",
                "description": "описание соответствия"
            }}
        ],
        "contradictions": [
            {{
                "doc_reference": "ссылка на пункт/раздел документа или 'отсутствует'",
                "rasp_reference": "ссылка на пункт Плана мероприятий",
                "description": "описание расхождения или отсутствующего элемента",
                "recommendation": "рекомендация по устранению"
            }}
        ]
    }},

    "conclusion": {{
        "status": "соответствует" | "частично соответствует" | "требует доработки",
        "summary": "итоговое заключение (2-3 предложения)",
        "priority_recommendations": ["список приоритетных рекомендаций по доработке"]
    }}
}}

ВАЖНО:
- Если соответствий или противоречий нет, оставь пустой массив []
- Будь конкретен в ссылках на пункты документов
- Учитывай специфику образовательной организации
- Отвечай ТОЛЬКО валидным JSON без дополнительного текста
- ОБЯЗАТЕЛЬНО указывай конкретные пункты/разделы/статьи анализируемого документа
- Рекомендации должны быть привязаны к конкретному содержанию документа, а не общими фразами
"""

# ==================== ПРОМПТ ВАЛИДАЦИИ ВХОДНЫХ ДАННЫХ ====================

VALIDATION_PROMPT = """Ты — эксперт по классификации документов.

Определи, является ли документ НОРМАТИВНО-ПРАВОВЫМ АКТОМ (НПА).

=== ДОКУМЕНТ ===
Название: {doc_name}

Текст (первые 3000 символов):
{doc_text}

=== ЧТО ТАКОЕ НПА ===
Нормативно-правовой акт — это официальный документ, который:
- Устанавливает, изменяет или отменяет правовые нормы
- Имеет обязательную юридическую силу
- Принят уполномоченным органом (государственным или организацией)

ПРИМЕРЫ НПА:
- Положения, регламенты, порядки, правила
- Приказы, распоряжения (нормативного характера)
- Инструкции, стандарты
- Уставы, кодексы
- Локальные акты организаций

НЕ ЯВЛЯЕТСЯ НПА:
- Художественная литература, книги, романы
- Научные статьи, диссертации, рефераты
- Новости, статьи СМИ, блоги
- Личная переписка, письма
- Рекламные материалы
- Договоры, контракты (индивидуальные акты)
- Справки, выписки, отчёты
- Презентации, методички

=== ОТВЕТ ===
Верни JSON строго в формате:
{{
    "is_npa": true/false,
    "document_type": "тип документа",
    "reason": "краткое обоснование (1 предложение)"
}}

Отвечай ТОЛЬКО валидным JSON."""

# ==================== ПРОМПТ ВАЛИДАЦИИ ВЫХОДНЫХ ДАННЫХ ====================

OUTPUT_VALIDATION_PROMPT = """Ты — контролёр качества экспертных заключений.

Проверь качество экспертного анализа документа. Анализ должен быть конкретным и привязанным к содержанию документа.

=== НАЗВАНИЕ ДОКУМЕНТА ===
{doc_name}

=== КРАТКОЕ СОДЕРЖАНИЕ ДОКУМЕНТА ===
{doc_summary}

=== РЕЗУЛЬТАТ АНАЛИЗА ===
{analysis_json}

=== КРИТЕРИИ КАЧЕСТВА ===

ХОРОШИЙ анализ:
- Ссылается на КОНКРЕТНЫЕ пункты, разделы, статьи анализируемого документа (например: "п. 3.2", "раздел 5", "ст. 12")
- Рекомендации содержат конкретные предложения по изменению текста
- Описания соответствий/расхождений упоминают специфику документа

ПЛОХОЙ анализ:
- Общие фразы без конкретики ("документ соответствует принципам", "рекомендуется дополнить")
- Отсутствие ссылок на пункты анализируемого документа
- Шаблонные рекомендации, не связанные с содержанием

=== ОТВЕТ ===
Верни JSON строго в формате:
{{
    "quality": "good/acceptable/poor",
    "has_specific_references": true/false,
    "has_specific_recommendations": true/false,
    "issues": ["список проблем если есть"],
    "verdict": "краткий вердикт (1 предложение)"
}}

Отвечай ТОЛЬКО валидным JSON."""

# ==================== ФУНКЦИИ ПАРСИНГА ====================

def extract_text_from_pdf(file_path: Path) -> str:
    try:
        doc = fitz.open(str(file_path))
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Ошибка чтения PDF: {str(e)}")


def extract_text_from_docx(file_path: Path) -> str:
    try:
        doc = Document(str(file_path))
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Ошибка чтения DOCX: {str(e)}")


def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        return extract_text_from_docx(file_path)
    elif suffix == ".txt":
        return file_path.read_text(encoding="utf-8")
    else:
        raise HTTPException(status_code=400, detail=f"Неподдерживаемый формат: {suffix}. Используйте PDF, DOCX или TXT")


# ==================== LLM ИНТЕГРАЦИЯ ====================

async def call_llm(prompt: str, max_tokens: int = 2048) -> str:
    """Базовый вызов LLM"""
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://krechet.space",
        "X-Title": "RGGU Expertise Service"
    }

    payload = {
        "model": "anthropic/claude-3.5-sonnet",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens,
        "temperature": 0.1
    }

    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.post(OPENROUTER_URL, json=payload, headers=headers)
        if response.status_code != 200:
            raise HTTPException(status_code=500, detail=f"Ошибка API: {response.status_code}")

        result = response.json()
        content = result["choices"][0]["message"]["content"]

        # Извлекаем JSON из ответа
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0]
        elif "```" in content:
            content = content.split("```")[1].split("```")[0]

        return content.strip()


async def validate_input_document(doc_name: str, doc_text: str) -> dict:
    """Проверка: является ли документ нормативно-правовым актом"""
    prompt = VALIDATION_PROMPT.format(
        doc_name=doc_name,
        doc_text=doc_text[:3000]
    )

    try:
        content = await call_llm(prompt, max_tokens=256)
        return json.loads(content)
    except (json.JSONDecodeError, Exception) as e:
        # При ошибке валидации пропускаем документ на анализ (fail-open)
        return {"is_npa": True, "reason": "Ошибка валидации, документ допущен к анализу"}


async def validate_output_quality(doc_name: str, doc_summary: str, analysis: dict) -> dict:
    """Проверка качества выходного анализа"""
    prompt = OUTPUT_VALIDATION_PROMPT.format(
        doc_name=doc_name,
        doc_summary=doc_summary[:1000] if doc_summary else "Не указано",
        analysis_json=json.dumps(analysis, ensure_ascii=False, indent=2)[:3000]
    )

    try:
        content = await call_llm(prompt, max_tokens=512)
        return json.loads(content)
    except (json.JSONDecodeError, Exception):
        return {"quality": "acceptable", "verdict": "Проверка качества не выполнена"}


async def analyze_document(doc_name: str, doc_text: str) -> dict:
    prompt = ANALYSIS_PROMPT.format(
        ukaz_text=UKAZ_809_TEXT,
        rasporyazhenie_text=RASPORYAZHENIE_1734_TEXT,
        doc_name=doc_name,
        doc_text=doc_text[:50000]
    )

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://krechet.space",
        "X-Title": "RGGU Expertise Service"
    }

    payload = {
        "model": "anthropic/claude-3.5-sonnet",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 4096,
        "temperature": 0.1
    }

    try:
        async with httpx.AsyncClient(timeout=55.0) as client:
            response = await client.post(OPENROUTER_URL, json=payload, headers=headers)
            if response.status_code != 200:
                raise HTTPException(status_code=500, detail=f"Ошибка API: {response.status_code} - {response.text[:200]}")

            result = response.json()
            content = result["choices"][0]["message"]["content"]

            try:
                if "```json" in content:
                    content = content.split("```json")[1].split("```")[0]
                elif "```" in content:
                    content = content.split("```")[1].split("```")[0]
                return json.loads(content.strip())
            except json.JSONDecodeError as e:
                return {"document_name": doc_name, "raw_response": content, "parse_error": str(e)}
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Таймаут при обращении к AI. Попробуйте ещё раз.")


# ==================== ГЕНЕРАЦИЯ DOCX ====================

def create_expertise_docx(analysis: dict) -> io.BytesIO:
    doc = Document()
    title = doc.add_heading("ЭКСПЕРТНОЕ ЗАКЛЮЧЕНИЕ", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph(f"Наименование документа: {analysis.get('document_name', 'Не указано')}")
    doc.add_paragraph(f"Дата экспертизы: {datetime.now().strftime('%d.%m.%Y')}")

    if analysis.get('summary'):
        doc.add_paragraph()
        doc.add_paragraph(f"Краткое содержание: {analysis['summary']}")

    doc.add_heading("АНАЛИЗ НА СООТВЕТСТВИЕ УКАЗУ ПРЕЗИДЕНТА РФ ОТ 09.11.2022 № 809", level=1)
    ukaz = analysis.get('ukaz_809', {})

    doc.add_heading("Соответствует:", level=2)
    for m in ukaz.get('matches', []):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{m.get('doc_reference', '')}: ").bold = True
        p.add_run(f"{m.get('description', '')} ")
        p.add_run(f"(соотв. {m.get('ukaz_reference', '')})").italic = True
    if not ukaz.get('matches'):
        doc.add_paragraph("Явных соответствий не выявлено.", style='List Bullet')

    doc.add_heading("Расходится:", level=2)
    for c in ukaz.get('contradictions', []):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{c.get('doc_reference', '')}: ").bold = True
        p.add_run(f"{c.get('description', '')} ")
        p.add_run(f"(см. {c.get('ukaz_reference', '')})").italic = True
    if not ukaz.get('contradictions'):
        doc.add_paragraph("Расхождений не выявлено.", style='List Bullet')

    doc.add_heading("АНАЛИЗ НА СООТВЕТСТВИЕ РАСПОРЯЖЕНИЮ ПРАВИТЕЛЬСТВА РФ ОТ 01.07.2024 № 1734-р", level=1)
    rasp = analysis.get('rasporyazhenie_1734', {})

    doc.add_heading("Соответствует:", level=2)
    for m in rasp.get('matches', []):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{m.get('doc_reference', '')}: ").bold = True
        p.add_run(f"{m.get('description', '')} ")
        p.add_run(f"(соотв. {m.get('rasp_reference', '')})").italic = True
    if not rasp.get('matches'):
        doc.add_paragraph("Явных соответствий не выявлено.", style='List Bullet')

    doc.add_heading("Расходится:", level=2)
    for c in rasp.get('contradictions', []):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{c.get('doc_reference', '')}: ").bold = True
        p.add_run(f"{c.get('description', '')} ")
        p.add_run(f"(см. {c.get('rasp_reference', '')})").italic = True
    if not rasp.get('contradictions'):
        doc.add_paragraph("Расхождений не выявлено.", style='List Bullet')

    doc.add_heading("ИТОГОВОЕ ЗАКЛЮЧЕНИЕ", level=1)
    conclusion = analysis.get('conclusion', {})
    p = doc.add_paragraph()
    p.add_run("Статус: ").bold = True
    p.add_run(conclusion.get('status', 'не определён').upper())
    if conclusion.get('summary'):
        doc.add_paragraph(conclusion['summary'])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==================== HTML СТРАНИЦА ====================

HTML_PAGE = '''<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Экспертиза НПА РГГУ</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        .drag-over { border-color: #3b82f6 !important; background-color: #eff6ff !important; }
        .fade-in { animation: fadeIn 0.3s ease-in; }
        @keyframes fadeIn { from { opacity: 0; transform: translateY(-10px); } to { opacity: 1; transform: translateY(0); } }
        .spinner { border: 3px solid #f3f3f3; border-top: 3px solid #3b82f6; border-radius: 50%; width: 20px; height: 20px; animation: spin 1s linear infinite; display: inline-block; }
        @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
        .pulse { animation: pulse 2s infinite; }
        @keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.5; } }
        .queue-item { transition: all 0.3s ease; }
        .queue-item.processing { background: linear-gradient(90deg, #dbeafe, #eff6ff, #dbeafe); background-size: 200% 100%; animation: shimmer 1.5s infinite; }
        @keyframes shimmer { 0% { background-position: 200% 0; } 100% { background-position: -200% 0; } }
    </style>
</head>
<body class="bg-gray-50 min-h-screen">
    <div class="container mx-auto px-4 py-8 max-w-4xl">
        <div class="text-center mb-8">
            <h1 class="text-3xl font-bold text-gray-800 mb-2">Экспертиза локальных актов РГГУ</h1>
            <p class="text-gray-600">Анализ на соответствие Указу № 809 и Распоряжению № 1734-р</p>
        </div>

        <div class="bg-white rounded-xl shadow-lg p-6 mb-8">
            <div id="dropZone" class="border-2 border-dashed border-gray-300 rounded-lg p-12 text-center cursor-pointer transition-all hover:border-blue-400 hover:bg-blue-50">
                <div class="text-6xl mb-4">📁</div>
                <p class="text-lg text-gray-700 mb-2">Перетащите файлы сюда</p>
                <p class="text-gray-500 mb-4">или нажмите для выбора</p>
                <p class="text-sm text-gray-400">PDF, DOCX, TXT (до 20 файлов)</p>
                <input type="file" id="fileInput" multiple accept=".pdf,.docx,.txt" class="hidden">
            </div>
            <div id="fileList" class="mt-4 hidden">
                <h3 class="font-semibold text-gray-700 mb-2">Выбранные файлы:</h3>
                <ul id="fileListItems" class="space-y-2"></ul>
            </div>
            <button id="analyzeBtn" class="w-full mt-6 bg-blue-600 text-white py-3 px-6 rounded-lg font-semibold hover:bg-blue-700 transition-colors disabled:bg-gray-400 hidden">▶ Провести экспертизу</button>
        </div>

        <!-- Очередь обработки -->
        <div id="queueSection" class="bg-white rounded-xl shadow-lg p-6 mb-8 hidden">
            <div class="flex justify-between items-center mb-4">
                <h2 class="text-lg font-bold text-gray-800">📊 Прогресс анализа</h2>
                <div class="flex items-center gap-2">
                    <span id="progressCounter" class="text-sm font-medium text-blue-600 bg-blue-100 px-3 py-1 rounded-full">0 / 0</span>
                </div>
            </div>
            <div class="w-full bg-gray-200 rounded-full h-3 mb-4">
                <div id="progressBar" class="bg-blue-600 h-3 rounded-full transition-all duration-500" style="width: 0%"></div>
            </div>
            <div id="queueList" class="space-y-2 max-h-48 overflow-y-auto"></div>
        </div>

        <div id="resultsSection" class="hidden">
            <div class="flex justify-between items-center mb-4">
                <h2 class="text-xl font-bold text-gray-800">Результаты экспертизы</h2>
                <button id="exportAllBtn" class="bg-green-600 text-white py-2 px-4 rounded-lg font-semibold hover:bg-green-700 hidden">📥 Скачать все рекомендации (ZIP)</button>
            </div>
            <div id="resultsList" class="space-y-4"></div>
        </div>
    </div>

    <script>
        const dropZone = document.getElementById('dropZone');
        const fileInput = document.getElementById('fileInput');
        const fileList = document.getElementById('fileList');
        const fileListItems = document.getElementById('fileListItems');
        const analyzeBtn = document.getElementById('analyzeBtn');
        const queueSection = document.getElementById('queueSection');
        const queueList = document.getElementById('queueList');
        const progressCounter = document.getElementById('progressCounter');
        const progressBar = document.getElementById('progressBar');
        const resultsSection = document.getElementById('resultsSection');
        const resultsList = document.getElementById('resultsList');
        const exportAllBtn = document.getElementById('exportAllBtn');

        let selectedFiles = [];
        let currentSessionId = null;
        let allResults = [];
        let resultIndex = 0;

        dropZone.addEventListener('click', () => fileInput.click());
        dropZone.addEventListener('dragover', (e) => { e.preventDefault(); dropZone.classList.add('drag-over'); });
        dropZone.addEventListener('dragleave', () => dropZone.classList.remove('drag-over'));
        dropZone.addEventListener('drop', (e) => { e.preventDefault(); dropZone.classList.remove('drag-over'); handleFiles(e.dataTransfer.files); });
        fileInput.addEventListener('change', (e) => handleFiles(e.target.files));

        function handleFiles(files) {
            selectedFiles = Array.from(files).slice(0, 20);
            if (selectedFiles.length === 0) { fileList.classList.add('hidden'); analyzeBtn.classList.add('hidden'); return; }
            fileList.classList.remove('hidden');
            analyzeBtn.classList.remove('hidden');
            fileListItems.innerHTML = selectedFiles.map((file, i) => `<li class="flex items-center justify-between bg-gray-50 rounded-lg px-4 py-2"><span>📄 ${file.name}</span><button onclick="removeFile(${i})" class="text-red-500 hover:text-red-700">✕</button></li>`).join('');
        }

        function removeFile(index) { selectedFiles.splice(index, 1); handleFiles(selectedFiles); }

        function updateQueue(files, currentIndex, statuses) {
            queueList.innerHTML = files.map((file, i) => {
                let statusIcon, statusClass, statusText;
                if (statuses[i] === 'done') {
                    statusIcon = '✅';
                    statusClass = 'bg-green-50 border-green-200';
                    statusText = 'Готово';
                } else if (statuses[i] === 'error') {
                    statusIcon = '❌';
                    statusClass = 'bg-red-50 border-red-200';
                    statusText = 'Ошибка';
                } else if (i === currentIndex) {
                    statusIcon = '<div class="spinner"></div>';
                    statusClass = 'processing border-blue-300';
                    statusText = 'Анализирую...';
                } else {
                    statusIcon = '⏳';
                    statusClass = 'bg-gray-50 border-gray-200';
                    statusText = 'В очереди';
                }
                return `<div class="queue-item flex items-center justify-between px-4 py-3 rounded-lg border ${statusClass}">
                    <div class="flex items-center gap-3">
                        <span class="w-6 h-6 flex items-center justify-center">${statusIcon}</span>
                        <span class="text-sm font-medium text-gray-700 truncate max-w-xs">${file.name}</span>
                    </div>
                    <span class="text-xs text-gray-500">${statusText}</span>
                </div>`;
            }).join('');
        }

        async function analyzeSequentially() {
            if (selectedFiles.length === 0) return;

            analyzeBtn.disabled = true;
            analyzeBtn.classList.add('hidden');
            fileList.classList.add('hidden');
            dropZone.classList.add('hidden');
            queueSection.classList.remove('hidden');
            resultsSection.classList.remove('hidden');
            exportAllBtn.classList.add('hidden');
            resultsList.innerHTML = '';

            const totalFiles = selectedFiles.length;
            const statuses = new Array(totalFiles).fill('pending');
            allResults = [];
            resultIndex = 0;

            // Создаем session_id на клиенте
            currentSessionId = crypto.randomUUID();

            progressCounter.textContent = `0 / ${totalFiles}`;
            progressBar.style.width = '0%';
            updateQueue(selectedFiles, 0, statuses);

            for (let i = 0; i < totalFiles; i++) {
                const file = selectedFiles[i];
                updateQueue(selectedFiles, i, statuses);

                const formData = new FormData();
                formData.append('file', file);
                formData.append('session_id', currentSessionId);
                formData.append('index', i.toString());

                try {
                    const response = await fetch('/nparggu/api/analyze-single', { method: 'POST', body: formData });
                    if (!response.ok) {
                        const errorData = await response.json().catch(() => ({ detail: 'Ошибка сервера' }));
                        throw new Error(errorData.detail || 'Ошибка сервера');
                    }
                    const result = await response.json();
                    result.filename = file.name;
                    allResults.push(result);
                    statuses[i] = result.error ? 'error' : 'done';

                    // Сразу добавляем результат в список
                    addResultCard(result, resultIndex);
                    resultIndex++;
                } catch (error) {
                    statuses[i] = 'error';
                    allResults.push({ filename: file.name, error: error.message });
                    addResultCard({ filename: file.name, error: error.message }, resultIndex);
                    resultIndex++;
                }

                // Обновляем прогресс
                const completed = i + 1;
                const remaining = totalFiles - completed;
                progressCounter.textContent = `${completed} / ${totalFiles}`;
                progressBar.style.width = `${(completed / totalFiles) * 100}%`;
                updateQueue(selectedFiles, i + 1, statuses);
            }

            // Всё готово
            queueSection.innerHTML = `
                <div class="text-center py-4">
                    <div class="text-4xl mb-2">🎉</div>
                    <p class="text-lg font-semibold text-green-600">Анализ завершён!</p>
                    <p class="text-sm text-gray-500">Обработано документов: ${totalFiles}</p>
                </div>
            `;

            exportAllBtn.classList.remove('hidden');

            // Возвращаем возможность загрузить новые файлы
            setTimeout(() => {
                dropZone.classList.remove('hidden');
                analyzeBtn.disabled = false;
                selectedFiles = [];
            }, 1000);
        }

        analyzeBtn.addEventListener('click', analyzeSequentially);

        function addResultCard(result, index) {
            const card = document.createElement('div');
            card.className = 'fade-in';
            card.innerHTML = renderResultCard(result, index);
            resultsList.appendChild(card);
            card.scrollIntoView({ behavior: 'smooth', block: 'nearest' });
        }

        function renderResultCard(result, index) {
            if (result.error) {
                const validationDetails = result.validation_details;
                const isValidationError = result.error.includes('нормативно-правовой акт');
                return `<div class="bg-white rounded-xl shadow-lg p-6 border-l-4 border-${isValidationError ? 'orange' : 'red'}-500">
                    <div class="flex items-center gap-2">
                        <span class="text-xl">${isValidationError ? '⚠️' : '❌'}</span>
                        <span class="font-semibold">${result.filename}</span>
                    </div>
                    <p class="text-${isValidationError ? 'orange' : 'red'}-600 mt-2">${result.error}</p>
                    ${validationDetails ? `<p class="text-sm text-gray-500 mt-1">${validationDetails.document_type || ''}${validationDetails.reason ? ': ' + validationDetails.reason : ''}</p>` : ''}
                </div>`;
            }

            const ukaz = result.ukaz_809 || {};
            const rasp = result.rasporyazhenie_1734 || {};
            const conclusion = result.conclusion || {};
            const qualityCheck = result.quality_check || {};
            const inputValidation = result.input_validation || {};
            const statusColor = conclusion.status === 'соответствует' ? 'green' : conclusion.status === 'частично соответствует' ? 'yellow' : 'red';
            const qualityColor = qualityCheck.quality === 'good' ? 'green' : qualityCheck.quality === 'acceptable' ? 'yellow' : 'red';

            const renderList = (items, type) => {
                if (!items || items.length === 0) return '<p class="text-gray-400 text-sm italic">Не выявлено</p>';
                return items.map(item => {
                    const ref = type === 'ukaz' ? item.ukaz_reference : item.rasp_reference;
                    return `<div class="mb-2 text-sm"><b>${item.doc_reference || ''}</b>: ${item.description || ''} <span class="text-gray-500">(${ref || ''})</span>${item.recommendation ? '<br><span class="text-blue-600">→ ' + item.recommendation + '</span>' : ''}</div>`;
                }).join('');
            };

            return `<div class="bg-white rounded-xl shadow-lg overflow-hidden border-l-4 border-${statusColor}-500">
                <div class="p-6">
                    <div class="flex justify-between items-start mb-3">
                        <div>
                            <h3 class="font-semibold text-lg">📄 ${result.filename || result.document_name}</h3>
                            <p class="text-sm text-gray-500 mt-1">${result.summary || ''}</p>
                            ${inputValidation.document_type ? `<p class="text-xs text-gray-400 mt-1">📋 ${inputValidation.document_type}${inputValidation.organization ? ' | ' + inputValidation.organization : ''}</p>` : ''}
                        </div>
                        <div class="flex flex-col items-end gap-1">
                            <span class="px-3 py-1 rounded-full text-sm font-medium bg-${statusColor}-100 text-${statusColor}-800">${conclusion.status || 'Анализ завершён'}</span>
                            ${qualityCheck.quality ? `<span class="px-2 py-0.5 rounded text-xs bg-${qualityColor}-50 text-${qualityColor}-700" title="${qualityCheck.verdict || ''}">🔍 ${qualityCheck.quality === 'good' ? 'Качество: высокое' : qualityCheck.quality === 'acceptable' ? 'Качество: приемлемое' : 'Качество: проверьте'}</span>` : ''}
                        </div>
                    </div>
                    ${conclusion.quality_warning ? `<div class="bg-orange-50 border border-orange-200 rounded-lg p-2 mb-3"><p class="text-xs text-orange-700">⚠️ ${conclusion.quality_warning}</p></div>` : ''}
                    <div class="grid grid-cols-2 gap-4 mb-4">
                        <div class="bg-gray-50 rounded-lg p-3">
                            <p class="text-sm font-medium text-gray-700">Указ № 809</p>
                            <p class="text-sm mt-1"><span class="text-green-600">✅ ${(ukaz.matches||[]).length} совпад.</span> | <span class="text-red-600">❌ ${(ukaz.contradictions||[]).length} расхожд.</span></p>
                        </div>
                        <div class="bg-gray-50 rounded-lg p-3">
                            <p class="text-sm font-medium text-gray-700">Распоряжение № 1734-р</p>
                            <p class="text-sm mt-1"><span class="text-green-600">✅ ${(rasp.matches||[]).length} совпад.</span> | <span class="text-red-600">❌ ${(rasp.contradictions||[]).length} расхожд.</span></p>
                        </div>
                    </div>
                    ${conclusion.summary ? `<div class="bg-blue-50 rounded-lg p-3 mb-4"><p class="text-sm font-medium text-blue-800">📝 Итог:</p><p class="text-sm text-blue-700 mt-1">${conclusion.summary}</p></div>` : ''}
                    <div class="flex justify-between items-center">
                        <button onclick="toggleDetails(${index})" class="text-blue-600 hover:text-blue-800 text-sm font-medium" id="toggleBtn${index}">▼ Подробнее</button>
                        <button onclick="downloadSingle(${index})" class="bg-blue-600 text-white px-4 py-2 rounded-lg text-sm hover:bg-blue-700">📥 Скачать DOCX</button>
                    </div>
                </div>
                <div id="details${index}" class="hidden border-t bg-gray-50 p-6">
                    <div class="grid md:grid-cols-2 gap-6">
                        <div>
                            <h4 class="font-bold text-gray-800 mb-3">📜 Указ № 809</h4>
                            <div class="mb-4"><p class="text-green-700 font-medium mb-2">✅ Соответствует:</p>${renderList(ukaz.matches, 'ukaz')}</div>
                            <div><p class="text-red-700 font-medium mb-2">❌ Расходится:</p>${renderList(ukaz.contradictions, 'ukaz')}</div>
                        </div>
                        <div>
                            <h4 class="font-bold text-gray-800 mb-3">📋 Распоряжение № 1734-р</h4>
                            <div class="mb-4"><p class="text-green-700 font-medium mb-2">✅ Соответствует:</p>${renderList(rasp.matches, 'rasp')}</div>
                            <div><p class="text-red-700 font-medium mb-2">❌ Расходится:</p>${renderList(rasp.contradictions, 'rasp')}</div>
                        </div>
                    </div>
                    ${conclusion.priority_recommendations && conclusion.priority_recommendations.length > 0 ?
                        `<div class="mt-6 bg-yellow-50 rounded-lg p-4"><p class="font-medium text-yellow-800 mb-2">⚡ Приоритетные рекомендации:</p><ul class="list-disc list-inside text-sm text-yellow-700">${conclusion.priority_recommendations.map(r => '<li>' + r + '</li>').join('')}</ul></div>` : ''}
                    ${qualityCheck.verdict ? `<div class="mt-4 text-xs text-gray-500 border-t pt-3"><b>🔍 Проверка качества:</b> ${qualityCheck.verdict}</div>` : ''}
                </div>
            </div>`;
        }

        function toggleDetails(index) {
            const details = document.getElementById('details' + index);
            const btn = document.getElementById('toggleBtn' + index);
            if (details.classList.contains('hidden')) {
                details.classList.remove('hidden');
                btn.textContent = '▲ Скрыть';
            } else {
                details.classList.add('hidden');
                btn.textContent = '▼ Подробнее';
            }
        }

        function downloadSingle(index) { if (currentSessionId) window.location.href = `/nparggu/api/export/${currentSessionId}/${index}`; }
        exportAllBtn.addEventListener('click', () => { if (currentSessionId) window.location.href = `/nparggu/api/export-all/${currentSessionId}`; });
    </script>
</body>
</html>'''


# ==================== API ЭНДПОИНТЫ ====================

@app.get("/", response_class=HTMLResponse)
@app.get("/nparggu", response_class=HTMLResponse)
@app.get("/nparggu/", response_class=HTMLResponse)
async def index():
    return HTML_PAGE


@app.post("/nparggu/api/analyze")
async def analyze_files(files: List[UploadFile] = File(...)):
    if len(files) > 20:
        raise HTTPException(status_code=400, detail="Максимум 20 файлов")

    session_id = str(uuid.uuid4())
    results = []

    for file in files:
        safe_filename = re.sub(r'[^\w\.\-]', '_', file.filename)
        file_path = UPLOADS_DIR / f"{session_id}_{safe_filename}"
        content = await file.read()
        file_path.write_bytes(content)

        try:
            text = extract_text(file_path)
            if not text.strip():
                results.append({"filename": file.filename, "error": "Не удалось извлечь текст"})
                continue
            analysis = await analyze_document(file.filename, text)
            analysis["filename"] = file.filename
            results.append(analysis)
        except Exception as e:
            results.append({"filename": file.filename, "error": str(e)})
        finally:
            if file_path.exists():
                file_path.unlink()

    results_store[session_id] = results
    return {"session_id": session_id, "results": results}


@app.post("/nparggu/api/analyze-single")
async def analyze_single_file(
    file: UploadFile = File(...),
    session_id: str = Form(None),
    index: int = Form(0)
):
    """Анализ одного файла с сохранением в сессию"""
    if not session_id:
        session_id = str(uuid.uuid4())

    safe_filename = re.sub(r'[^\w\.\-]', '_', file.filename)
    file_path = UPLOADS_DIR / f"{session_id}_{safe_filename}"
    content = await file.read()
    file_path.write_bytes(content)

    result = {}
    try:
        text = extract_text(file_path)
        if not text.strip():
            result = {"filename": file.filename, "error": "Не удалось извлечь текст из документа"}
        else:
            # ========== КОНТУР 1: Проверка входных данных ==========
            validation = await validate_input_document(file.filename, text)

            if not validation.get("is_npa", True):
                result = {
                    "filename": file.filename,
                    "error": "Это не нормативно-правовой акт, обратитесь к Михаилу.",
                    "validation_details": {
                        "document_type": validation.get("document_type", "Не НПА"),
                        "reason": validation.get("reason", "")
                    }
                }
            else:
                # Проводим анализ
                analysis = await analyze_document(file.filename, text)
                analysis["filename"] = file.filename

                # ========== КОНТУР 2: Проверка качества выходных данных ==========
                quality_check = await validate_output_quality(
                    file.filename,
                    analysis.get("summary", ""),
                    analysis
                )

                # Добавляем информацию о качестве анализа
                analysis["quality_check"] = {
                    "quality": quality_check.get("quality", "unknown"),
                    "verdict": quality_check.get("verdict", ""),
                    "has_specific_references": quality_check.get("has_specific_references", None)
                }

                # Добавляем информацию о валидации
                analysis["input_validation"] = {
                    "document_type": validation.get("document_type", "НПА"),
                    "is_npa": True
                }

                # Если качество плохое, добавляем предупреждение
                if quality_check.get("quality") == "poor":
                    if "conclusion" not in analysis:
                        analysis["conclusion"] = {}
                    analysis["conclusion"]["quality_warning"] = "Внимание: анализ может содержать общие рекомендации. Рекомендуется ручная проверка."

                result = analysis

    except Exception as e:
        result = {"filename": file.filename, "error": str(e)}
    finally:
        if file_path.exists():
            file_path.unlink()

    # Сохраняем результат в сессию
    if session_id not in results_store:
        results_store[session_id] = []

    # Убедимся, что список достаточной длины
    while len(results_store[session_id]) <= index:
        results_store[session_id].append(None)

    results_store[session_id][index] = result

    return result


@app.get("/nparggu/api/export/{session_id}/{index}")
async def export_single(session_id: str, index: int):
    if session_id not in results_store:
        raise HTTPException(status_code=404, detail="Сессия не найдена")
    results = results_store[session_id]
    if index < 0 or index >= len(results):
        raise HTTPException(status_code=404, detail="Документ не найден")
    analysis = results[index]
    if "error" in analysis:
        raise HTTPException(status_code=400, detail="Нет экспертизы")

    docx_buffer = create_expertise_docx(analysis)
    filename_ascii = f"Expertiza_{index}.docx"
    filename_utf8 = quote(f"Экспертиза_{index}.docx")

    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename_ascii}"; filename*=UTF-8\'\'{filename_utf8}'}
    )


@app.get("/nparggu/api/export-all/{session_id}")
async def export_all(session_id: str):
    if session_id not in results_store:
        raise HTTPException(status_code=404, detail="Сессия не найдена")
    results = results_store[session_id]

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for i, analysis in enumerate(results):
            if "error" not in analysis:
                docx_buffer = create_expertise_docx(analysis)
                zf.writestr(f"Expertiza_{i}.docx", docx_buffer.getvalue())
    zip_buffer.seek(0)

    filename = f"Expertiza_RGGU_{session_id[:8]}.zip"
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )
